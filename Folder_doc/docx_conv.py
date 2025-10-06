import os
import subprocess
from docx import Document
from bs4 import BeautifulSoup  # kept in case you still want to clean HTML
import pandas as pd

LIBREOFFICE_PATH = r"C:\Program Files\LibreOffice\program\soffice.exe"


def convert_docx_to_pdf(input_docx, output_pdf):
    """Convert DOCX to PDF using LibreOffice CLI."""
    if not os.path.exists(input_docx):
        print("DOCX file not found.")
        return

    try:
        subprocess.run([
            LIBREOFFICE_PATH,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", os.path.dirname(output_pdf),
            input_docx
        ], check=True)
        print(f"Converted DOCX to PDF: {output_pdf}")
    except FileNotFoundError:
        print("LibreOffice not found. Please install or update LIBREOFFICE_PATH.")
    except subprocess.CalledProcessError as e:
        print(f"LibreOffice PDF conversion failed: {e}")

def convert_docx_to_txt(input_docx, output_txt):
    """Convert DOCX to plain text."""
    if not os.path.exists(input_docx):
        print("DOCX file not found.")
        return

    doc = Document(input_docx)
    text = "\n".join(para.text for para in doc.paragraphs)

    with open(output_txt, "w", encoding="utf-8") as f:
        f.write(text)

    print(f"Converted DOCX to TXT: {output_txt}")

def convert_docx_to_html(input_docx, output_html):
    """Convert DOCX to HTML using LibreOffice CLI."""
    if not os.path.exists(input_docx):
        print("DOCX file not found.")
        return

    try:
        subprocess.run([
            LIBREOFFICE_PATH,
            "--headless",
            "--convert-to", "html",
            "--outdir", os.path.dirname(output_html),
            input_docx
        ], check=True)
        print(f"Converted DOCX to HTML: {output_html}")
    except FileNotFoundError:
        print("LibreOffice not found. Please install or update LIBREOFFICE_PATH.")
    except subprocess.CalledProcessError as e:
        print(f"LibreOffice HTML conversion failed: {e}")

def convert_docx_to_odt(input_docx, output_odt):
    """Convert DOCX to ODT using LibreOffice CLI."""
    if not os.path.exists(input_docx):
        print("DOCX file not found.")
        return

    try:
        subprocess.run([
            LIBREOFFICE_PATH,
            "--headless",
            "--convert-to", "odt",
            "--outdir", os.path.dirname(output_odt),
            input_docx
        ], check=True)
        print(f"Converted DOCX to ODT: {output_odt}")
    except FileNotFoundError:
        print("LibreOffice not found. Please install or update LIBREOFFICE_PATH.")
    except subprocess.CalledProcessError as e:
        print(f"LibreOffice ODT conversion failed: {e}")

def convert_docx_to_doc(input_docx, output_doc):
    """Convert DOCX to DOC (old Word format) using LibreOffice."""
    if not os.path.exists(input_docx):
        print("DOCX file not found.")
        return

    try:
        subprocess.run([
            LIBREOFFICE_PATH,
            "--headless",
            "--convert-to", "doc:MS Word 97",
            "--outdir", os.path.dirname(output_doc),
            input_docx
        ], check=True)
        print(f"Converted DOCX to DOC: {output_doc}")
    except FileNotFoundError:
        print("LibreOffice not found. Please install or update LIBREOFFICE_PATH.")
    except subprocess.CalledProcessError as e:
        print(f"LibreOffice DOC conversion failed: {e}")

def convert_docx_to_csv(input_docx, output_csv):
    """Extract first table from DOCX and save as CSV."""
    if not os.path.exists(input_docx):
        print("DOCX file not found.")
        return

    doc = Document(input_docx)
    if not doc.tables:
        pd.DataFrame([["No tables found"]]).to_csv(output_csv, index=False)
        print(f"No tables found in DOCX. Created CSV with notice: {output_csv}")
        return

    table = doc.tables[0]
    data = [[cell.text.strip() for cell in row.cells] for row in table.rows]

    pd.DataFrame(data).to_csv(output_csv, index=False)
    print(f"Converted DOCX table to CSV: {output_csv}")
